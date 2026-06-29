"""
Build a comprehensive, plain-language Word (.docx) report of the PDAC ST + Foundation-Model
project. Walks the pipeline STAGE BY STAGE and embeds the real diagnostic output images from
each Outputs/ subfolder (not just the presentation figures), each with an easy-words explanation.

Run:  "C:/Users/datai/anaconda3/envs/tcga/python.exe" presentation/make_word_report.py
Output: Outputs/PDAC_Project_Report.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
PROJ = os.path.dirname(ROOT)
OUT = os.path.join(PROJ, "Outputs")
DOCX = os.path.join(OUT, "PDAC_Project_Report.docx")

BLUE = RGBColor(0x2C, 0x5F, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BLUE

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def callout(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = BLUE
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

def img(relpath, width=6.0, cap=None):
    path = os.path.join(PROJ, relpath)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        r = p.add_run(f"[missing image: {relpath}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        caption(cap)

# =====================================================================================
# TITLE PAGE
# =====================================================================================
t = doc.add_heading("Spatial-Transcriptomics-Guided Histopathology\nfor Metastatic-Propensity "
                    "Scoring in Pancreatic Cancer", level=0)
for r in t.runs:
    r.font.color.rgb = BLUE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("A stage-by-stage methods & results report, in plain language")
rs.italic = True
rs.font.size = Pt(13)
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("PDAC primary tumour -> liver metastasis  |  6 samples, 18,859 spots  |  "
             "H&E + spatial transcriptomics\n").font.size = Pt(11)

doc.add_paragraph()
callout("One-line summary")
body("We tried to flag metastasis-prone tumour regions from a routine H&E slide alone, using "
     "spatial transcriptomics as a teacher during training. The intra-tumour EMT \"leaving "
     "program\" we target is a real, externally validated biological signal -- but from H&E "
     "alone, on patients the model has never seen, the model currently recovers only how much "
     "tumour is present, not the fine confound-free metastatic state. The multimodal bridge added "
     "nothing over a standard image model. This is a rigorous, confound-audited result, and the "
     "report shows every diagnostic that led to it.")

body("How to read this report: it follows the pipeline in the order we actually ran it "
     "(Stage 0 -> Stage 4). Each stage states what we did, shows its real output figures from the "
     "Outputs/ folder, and explains in plain words what the figure tells us -- including the "
     "honest negatives. A glossary of terms is at the end.")

# =====================================================================================
# 1. BACKGROUND
# =====================================================================================
h1("1. Background and goal")
body("Pancreatic cancer (PDAC) is lethal largely because it spreads early, most often to the "
     "liver, often before diagnosis. The clue to \"which tumour regions are preparing to spread\" "
     "lives in the tumour's gene activity, which a normal H&E slide does not show, and which is "
     "expensive to measure routinely.")
callout("The idea: let gene data teach during training, then work from the cheap H&E alone.")
body("Like training a medical student beside a senior expert: during training the student sees "
     "the slide AND the expert's molecular read-out; afterwards the student must make the call "
     "from the slide alone. Hard constraint from Dr. Ashiq: spatial transcriptomics may be used "
     "to train, but inference must use H&E only.")

# =====================================================================================
# 2. DATA
# =====================================================================================
h1("2. The data and how we represent it")
img("Outputs/presentation_figures/fig1_cohort_overview.png", 6.3,
    "Figure 2.1 -- The cohort. 6 tissues (4 primary pancreatic tumours, 2 liver metastases), "
    "18,859 spots; each spot carries H&E + gene expression + cell composition.")
body("We have 6 samples: Primary Tumour (PT) T1, T3, T4, T11 and Hepatic Metastasis (HM) HM11, "
     "HM13. Only patient 11 has BOTH a primary (T11) and its matched liver metastasis (HM11). "
     "Each tissue spot (~55 micrometres) carries three views:")
bullet("H&E image patch (224x224 px) -> turned into a numeric fingerprint by a pathology "
       "foundation model (we use UNI2-h; also tested CONCH, H-Optimus, ResNet18).")
bullet("Gene expression of ~17,900 genes -> compressed into a clean 50-number summary by scVI.")
bullet("Cell composition -> proportions of 15 cell types (tumour, hepatocytes/liver, CAFs, "
       "immune cells, ...) estimated by RCTD.")

h2("How the image models 'see' the patches")
body("Before training anything, we checked how different foundation models organise the H&E "
     "patches. Each dot below is one patch, placed by similarity (t-SNE). Patches that look alike "
     "land together. This is a sanity check that the image fingerprints carry real structure.")
img("Outputs/Plots/Image_Embeddings/ImageEmbeddings_uni2h_tSNE.png", 5.6,
    "Figure 2.2 -- UNI2-h image embeddings (our main backbone).")
img("Outputs/Plots/Image_Embeddings/ImageEmbeddings_hoptimus_tSNE.png", 5.6,
    "Figure 2.3 -- H-Optimus-1 image embeddings.")
img("Outputs/Plots/Image_Embeddings/ImageEmbeddings_conchv15_tSNE.png", 5.6,
    "Figure 2.4 -- CONCH v1.5 image embeddings.")
img("Outputs/Plots/Image_Embeddings/ImageEmbeddings_conchv1_tSNE.png", 5.6,
    "Figure 2.5 -- CONCH v1 image embeddings.")
img("Outputs/Plots/Image_Embeddings/ImageEmbeddings_resnet18_tSNE.png", 5.6,
    "Figure 2.6 -- ResNet18 (a plain non-pathology baseline) for comparison.")
body("Takeaway: the large pathology foundation models produce richer, more organised structure "
     "than the plain ResNet baseline -- which is why we build on them.")

# =====================================================================================
# 3. STAGE 0
# =====================================================================================
h1("3. Stage 0 -- The confound check (the trap we avoided)")
body("The obvious first idea is: compare primary tumours to liver metastases and call the "
     "difference 'metastatic-ness'. We proved this is a trap, BEFORE training anything.")
img("Outputs/stage0_confound/umap_scvi.png", 6.0,
    "Figure 3.1 -- Map of all spots by gene profile (scVI), coloured by sample / tissue type / "
    "cell composition. Spots separate largely by which tissue and cell mix they belong to.")
img("Outputs/stage0_confound/hepatocyte_by_sample.png", 5.6,
    "Figure 3.2 -- Liver-cell (hepatocyte) content per sample. Liver-met spots are ~15% liver "
    "cells; primary-tumour spots are ~0%.")
body("So a model trained on 'met vs primary' would mostly learn 'does this look like liver?' -- "
     "tissue location, not metastatic biology. The figures below quantify exactly that.")
img("Outputs/stage0_confound/delta_projection.png", 5.6,
    "Figure 3.3 -- Projecting spots onto the naive 'metastasis axis' (mean-HM minus mean-PT). "
    "It separates HM from PT, but for the wrong reason.")
img("Outputs/stage0_confound/proj_vs_hepatocyte.png", 5.6,
    "Figure 3.4 -- That 'metastasis axis' is 0.40-correlated with liver content across all spots, "
    "dropping to just 0.07 once we restrict to tumour-rich spots. It really was mostly liver.")
callout("Decision: abandon cross-patient 'met vs primary'. Define the target INSIDE the primary "
        "tumours, where there is no liver to confuse us.")
body("We also found that with only 2 metastasis patients, one cannot predict the other (0% "
     "recovered), while the 4 primaries generalise to each other (93-100%). The two mets simply "
     "do not share a transferable metastasis program. This is why Stage 0 is free insurance: a "
     "flashy model would have 'worked' here and quietly been a liver detector.")

# =====================================================================================
# 4. STAGE 1A
# =====================================================================================
h1("4. Stage 1A -- Defining the target: the 'leaving program'")
body("When a cancer cell prepares to detach and spread it undergoes EMT "
     "(epithelial-to-mesenchymal transition) -- it stops being a tidy, stuck-in-place cell and "
     "becomes a mobile, invasive one. We scored each primary-tumour spot for how strongly 19 "
     "canonical EMT / invasion genes are switched on. That is the 'leaving-program' score "
     "(13,578 PT spots).")

h2("Is the pattern spatially real?")
img("Outputs/stage1a_leaving_program/heatmap_IU_PDA_T1.png", 6.3, "Figure 4.1 -- T1: H&E vs leaving score.")
img("Outputs/stage1a_leaving_program/heatmap_IU_PDA_T3.png", 6.3, "Figure 4.2 -- T3: H&E vs leaving score.")
img("Outputs/stage1a_leaving_program/heatmap_IU_PDA_T4.png", 6.3, "Figure 4.3 -- T4: H&E vs leaving score.")
img("Outputs/stage1a_leaving_program/heatmap_IU_PDA_T11.png", 6.3, "Figure 4.4 -- T11: H&E vs leaving score.")
body("High-score spots form coherent zones, not random scatter (Moran's I 0.37-0.61, all highly "
     "significant) -- consistent with invasive fronts. This is real tissue geography.")

h2("Is it the right biology, and is it self-consistent?")
img("Outputs/stage1a_leaving_program/module_vs_pca.png", 5.4,
    "Figure 4.5 -- Our gene-checklist score agrees with an unsupervised data-driven axis (r=0.77) "
    "-- it is not an artefact of one particular method.")
img("Outputs/stage1a_leaving_program/core_vs_heldout.png", 5.4,
    "Figure 4.6 -- The score (built from one EMT gene set) agrees with a DIFFERENT EMT gene set "
    "we kept out (corr 0.46) -- it generalises beyond the genes used to build it.")
img("Outputs/stage1a_leaving_program/score_distribution.png", 5.4,
    "Figure 4.7 -- Distribution of leaving scores per sample.")

h2("The confound inside the target -- and how we handle it")
img("Outputs/stage1a_leaving_program/confound_bars.png", 5.4,
    "Figure 4.8 -- The score is 0.61-correlated with how much tumour is in the spot (abundance), "
    "and negatively with CAF/liver content. Part of the score is 'how much cancer', not 'EMT'.")
callout("We keep TWO targets from here on: raw (includes abundance) and resid (abundance removed "
        "-> confound-free, the honest headline, but weaker).")

# =====================================================================================
# 5. STAGE 1B
# =====================================================================================
h1("5. Stage 1B -- The patient-11 matched anchor (a weak/null check)")
body("Patient 11 is the only one with a matched primary (PT11) and liver met (HM11). We asked: "
     "do high-leaving PT11 spots actually resemble HM11? This is a corroboration check only.")
img("Outputs/stage1b_pt11_anchor/pt_vs_hm_separation.png", 5.4,
    "Figure 5.1 -- PT11 vs HM11 separation in gene space.")
img("Outputs/stage1b_pt11_anchor/heatmap_pt11_resemblance.png", 6.0,
    "Figure 5.2 -- Spatial map of 'HM11-resemblance' across PT11.")
img("Outputs/stage1b_pt11_anchor/leaving_vs_resemblance.png", 5.4,
    "Figure 5.3 -- Leaving score vs HM11-resemblance: weak (+0.17 for the confound-free score, "
    "essentially nothing for raw).")
img("Outputs/stage1b_pt11_anchor/driver_genes.png", 5.6,
    "Figure 5.4 -- What drives the axis: PT11-like = scar/collagen (CAF) genes; HM11-like = liver "
    "(ALB) + immune (immunoglobulin) genes. The axis is microenvironment, not metastasis biology.")
callout("Verdict: the convergence does not hold up. We lean ENTIRELY on the Stage-1A intra-PT "
        "leaving program and treat 1B as honest, weak corroboration only.")

# =====================================================================================
# 6. STAGE 2
# =====================================================================================
h1("6. Stage 2 -- Phase A: the multimodal 'bridge' (training)")
body("Here we tried to line up the three views of each spot (image, genes, cells) in a shared "
     "space using contrastive learning -- think of three translators forced to place the same "
     "spot in the same location, so the image translator secretly absorbs the biology. We tested "
     "4 backbones and added leakage controls (balanced sampling across slides; excluding "
     "spatial neighbours when measuring matching).")

h2("Training curves -- the model trains, but does it generalise?")
img("Outputs/stage2/UNI2h/stage2_UNI2h_outputs/stage2_training_curves.png", 6.0,
    "Figure 6.1 -- UNI2-h training curves.")
img("Outputs/stage2/UNI2h/stage2_UNI2h_ckpts/kfold_training_curve.png", 6.0,
    "Figure 6.2 -- UNI2-h per-fold (leave-one-sample-out) curves.")
img("Outputs/stage2/H-Optimus-1/stage2_outputs/stage2_training_curves.png", 6.0,
    "Figure 6.3 -- H-Optimus-1 training curves.")
img("Outputs/stage2/conch-v1.5/stage2_CONCH-V1.5_outputs/stage2_training_curves.png", 6.0,
    "Figure 6.4 -- CONCH v1.5 training curves.")
img("Outputs/stage2/conch-v1/stage2_outputs/stage2_training_curves.png", 6.0,
    "Figure 6.5 -- CONCH v1 training curves.")
callout("Honest result: on unseen patients, with neighbours excluded, the bridge's matching "
        "ability is basically chance. Earlier 'great' numbers were a spatial-autocorrelation "
        "illusion (neighbours look identical).")

h2("Comparing the four backbones")
img("Outputs/stage2/_model_comparison/radar_variantA.png", 5.4,
    "Figure 6.6 -- Four-backbone comparison (Variant A). The two large models (UNI2-h, "
    "H-Optimus) lead on the meaningful axes.")
img("Outputs/stage2/_model_comparison/bars_raw_resid_AB.png", 6.0,
    "Figure 6.7 -- Raw vs confound-free (resid) prediction for Variant A vs the metastasis-aware "
    "Variant B. The aux head (B) does not robustly help; resid sits at the noise floor.")
body("We carried UNI2-h (Variant A) into Stage 3 -- but the signal that the bridge was not "
     "earning its keep was already clear.")

# =====================================================================================
# 7. STAGE 3
# =====================================================================================
h1("7. Stage 3 -- Phase B: scoring from H&E, and the ceiling")
body("We trained the final predictor the honest way -- leave-one-patient-out -- and ran a "
     "decisive ablation: does the multimodal bridge beat just using the off-the-shelf image "
     "model?")
img("Outputs/stage3_phase_b/ablation_bars.png", 5.6,
    "Figure 7.1 -- Decisive ablation: the bridge (0.289) ties the plain frozen image model "
    "(0.270); a feed-forward variant is worse. The bridge adds essentially nothing.")
img("Outputs/stage3_phase_b/pred_vs_target.png", 5.6,
    "Figure 7.2 -- Predicted-from-H&E vs actual leaving score. H&E predicts the raw score "
    "moderately (rho ~0.29); the confound-free score is near the noise floor (~0.09).")
img("Outputs/stage3_phase_b/confound_bars.png", 5.6,
    "Figure 7.3 -- Confound audit: the raw prediction is 0.51-correlated with tumour amount (it "
    "is largely an abundance detector); the confound-free prediction is decorrelated from all "
    "confounders -- clean, but near-empty.")

h2("Seeing the ceiling on the tissue")
img("Outputs/stage3_phase_b/heatmap_IU_PDA_T1.png", 6.3, "Figure 7.4 -- T1: H&E | actual | predicted-from-H&E.")
img("Outputs/stage3_phase_b/heatmap_IU_PDA_T3.png", 6.3, "Figure 7.5 -- T3: H&E | actual | predicted-from-H&E.")
img("Outputs/stage3_phase_b/heatmap_IU_PDA_T4.png", 6.3, "Figure 7.6 -- T4: H&E | actual | predicted-from-H&E.")
img("Outputs/stage3_phase_b/heatmap_IU_PDA_T11.png", 6.3, "Figure 7.7 -- T11: H&E | actual | predicted-from-H&E.")
body("The prediction captures the broad high/low zones (which track tumour density) but smooths "
     "away the fine detail -- the picture behind the numbers.")
img("Outputs/stage3_phase_b/convergence_pt11.png", 5.4,
    "Figure 7.8 -- PT11 -> HM11 convergence is negative: high-predicted PT11 spots anti-resemble "
    "HM11 (an artefact of the abundance link). This closes the 'two axes agree' idea for good.")
callout("Decision: drop the bridge from deployment. Final predictor = direct regression from the "
        "frozen image model to the leaving score (gene data still only defined the target; never "
        "used at inference).")

# =====================================================================================
# 8. STAGE 4
# =====================================================================================
h1("8. Stage 4 -- Independent validation (the verdict)")
body("We pre-registered the tests and brought in external evidence: the source paper's own "
     "spot-by-spot scores for 29 signatures (including EMT), matched to our spots 100%.")
img("Outputs/stage4_validation/external_emt.png", 5.6,
    "Figure 8.1 -- Our target genuinely matches the independent paper's EMT signature (+0.205) -- "
    "the biology is real. But the confound-free EMT predicted from H&E is at the noise floor "
    "(+0.006) -- it cannot be read from the slide on unseen patients.")
img("Outputs/stage4_validation/fges_specificity.png", 5.6,
    "Figure 8.2 -- What the H&E prediction actually tracks: broad immune/proliferation/region "
    "signals, not a specific EMT read-out.")
img("Outputs/stage4_validation/test1_pt11.png", 5.6,
    "Figure 8.3 -- Patient-11 anchor test: fails -- high-scoring primary spots do not resemble "
    "the matched liver met (microenvironment, not metastasis).")
img("Outputs/stage4_validation/margin_enrichment.png", 5.6,
    "Figure 8.4 -- Invasive-margin test: the hypothesis fails -- scores are higher in the tumour "
    "interior than at the margin (abundance-driven, not invasive-front-specific).")
img("Outputs/stage4_validation/negative_control.png", 5.6,
    "Figure 8.5 -- Negative control: the confound-free score is properly decorrelated from "
    "confounders (good), BUT a trivial 'how much tumour' model (+0.48) beats the H&E model "
    "(+0.29). The abundance ceiling is not beaten.")
img("Outputs/presentation_figures/fig7_stage4_scorecard.png", 6.3,
    "Figure 8.6 -- Scorecard: the six pre-registered tests at a glance.")

# =====================================================================================
# 9. WORKED EXAMPLE
# =====================================================================================
h1("9. A worked example -- one result, traceable end to end")
img("Outputs/presentation_figures/fig_worked_example.png", 6.5,
    "Figure 9.1 -- Two real spots in tumour T3: where they sit, their actual H&E, and proof that "
    "high-score spots switch the EMT genes on -- with an independent paper's EMT score agreeing.")
body("This shows the score is not a black box: tissue location -> what it looks like -> which "
     "genes are on -> our score -> an outside ruler agrees. A pathologist can check every link.")

# =====================================================================================
# 10. CONCLUSIONS
# =====================================================================================
h1("10. Conclusions and what would push it further")
callout("The locked, honest headline")
body("The intra-primary EMT 'leaving program' is a real, externally validated transcriptomic "
     "signal; but from a routine H&E slide our model currently recovers only how much tumour is "
     "present, not the confound-free metastatic state, on unseen patients. The multimodal bridge "
     "adds nothing over a standard image model.")
h2("Why it's limited (the real reasons)")
bullet("Resolution: each spot (~55 um) mixes several cells, blurring the subtle within-tumour "
       "EMT signal.")
bullet("Tiny matched cohort: only 1 patient with both primary and metastasis; only 2 mets total.")
bullet("No ground-truth pathology annotation of invasive fronts to score against yet.")
h2("What would close the gap")
bullet("Pathologist annotation of invasive margins (the single most valuable missing piece).")
bullet("Higher-resolution platforms (Visium HD / Xenium) to un-blur the EMT signal.")
bullet("More matched primary-metastasis patients so a cross-patient program is learnable.")

# =====================================================================================
# 11. GLOSSARY
# =====================================================================================
h1("11. Glossary")
glossary = [
    ("H&E", "the standard pink/purple stained tissue slide pathologists read."),
    ("Spatial transcriptomics (Visium)", "measures gene activity at thousands of tiny spots over "
     "a slide, keeping their X/Y positions."),
    ("Spot", "one tiny tissue location (~55 um) with its own image patch + gene readout + cell mix."),
    ("Foundation model (UNI2-h, CONCH)", "a large image network pre-trained on millions of "
     "pathology images; turns a patch into a numeric fingerprint."),
    ("Embedding", "a list of numbers summarising an image patch or gene profile for comparison."),
    ("scVI", "compresses ~17,900 noisy gene measurements into a clean 50-number summary."),
    ("RCTD", "estimates the mix of cell types in each spot (tumour %, liver %, scar %, ...)."),
    ("EMT / 'leaving program'", "the biological shift a cancer cell makes to detach, invade, spread."),
    ("Confound", "a sneaky alternative explanation (e.g. 'looks like liver') that masquerades as "
     "the thing you care about."),
    ("Residualize", "mathematically subtract a confound so what's left is the clean signal."),
    ("Contrastive learning", "training that pulls matching things together and pushes non-matching "
     "things apart in a shared space."),
    ("Leave-one-patient-out (LOSO)", "train on some patients, test on a patient never seen -- the "
     "real test of generalisation."),
    ("Moran's I", "a clumpiness score; high means high-value spots cluster (real geography, not noise)."),
    ("Correlation (rho)", "how strongly two things move together: 0 = unrelated, ~1 = lockstep. "
     "+0.29 moderate, +0.09 near nothing."),
    ("Noise floor", "the 'signal' you'd get from pure randomness; being at it means no real signal."),
]
for term, defn in glossary:
    p = doc.add_paragraph()
    r = p.add_run(term + ":  ")
    r.bold = True
    p.add_run(defn)

doc.save(DOCX)
print("Saved:", DOCX)
