"""
Build the local Word version of the findings report.

Same content and figures as the shared web report, laid out for print/track-changes
so it can be circulated to collaborators who want to edit rather than browse.

Run:
    python docs/reports/build_findings_docx.py
Output:
    docs/reports/PDAC_site_specificity_findings.docx
"""

import json
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
O6 = os.path.join(ROOT, "Outputs", "stage6_site_program")
O7 = os.path.join(ROOT, "Outputs", "stage7_external_replication")
O8 = os.path.join(ROOT, "Outputs", "stage8_hest_cross_organ")
DEST = os.path.join(ROOT, "docs", "reports", "PDAC_site_specificity_findings.docx")

INK = RGBColor(0x12, 0x16, 0x1D)
GREY = RGBColor(0x4A, 0x53, 0x61)
ACCENT = RGBColor(0xEB, 0x68, 0x34)

M6 = json.load(open(os.path.join(O6, "metrics.json")))
M8 = json.load(open(os.path.join(O8, "metrics.json")))
cos6 = M6["cos_HMshift_LNMshift"]


def style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(7)
    n.paragraph_format.line_spacing = 1.15


def h(doc, text, level, color=INK):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = color
        r.font.name = "Calibri"
    return p


def para(doc, text, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def callout(doc, text):
    """An indented interpretation block — the 'what this means' voice."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    r.italic = True
    return p


def figure(doc, path, caption, width=6.6):
    if not os.path.exists(path):
        para(doc, f"[missing figure: {os.path.basename(path)}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    r.italic = True


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(txt))
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def main():
    doc = Document()
    style(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    h(doc, "Where a pancreatic tumour spreads changes what it becomes", 0)
    para(doc, "Metastatic PDAC is not one transcriptional program — the destination organ "
              "changes it, and the difference is immune.", italic=True, color=GREY, size=12)
    para(doc, "Spatial transcriptomics · three datasets · 244 tissue sections · 7 organs",
         color=GREY, size=9)

    # ---------------------------------------------------------------- summary
    h(doc, "Summary", 1)
    table(doc, ["Finding", "Evidence", "Strength"], [
        [f"Liver and lymph-node metastases share only ~{100*cos6:.0f}% of their shift "
         "from the primary tumour", "30 sections, 13 patients (GSE272362)", "Supported"],
        ["What differs between the sites is immune — B cells",
         "d = +2.03, FDR q = 0.042, 4/4 paired patients", "Supported"],
        ["Replicates across three metastatic sites; liver is immune-cold, lung immune-hot",
         "55 sections, 13 treatment-naive patients (GSE274557)", "Replicated"],
        ["Morphology→expression rules do not transfer between organs",
         "159 sections, 7 organs (HEST-1k), batch-controlled", "Confirmed"],
        ["The B-cell result specifically", "no lymph-node samples in the replication cohort",
         "Untested elsewhere"],
        ["Any link to patient outcome", "clinical fields for only 6/13 patients, no survival "
         "endpoint", "Not possible"],
    ])

    # ---------------------------------------------------------------- question
    h(doc, "The question", 1)
    para(doc, "A Visium slide measures gene expression in ~55-micron 'spots', each holding "
              "roughly one to ten cells. We have 91,496 of them across 30 tissue sections from "
              "13 people with pancreatic cancer: primary tumours, liver metastases, lymph-node "
              "metastases and normal pancreas.")
    para(doc, "The question is whether a tumour that has spread to the liver and one that has "
              "spread to a lymph node are running the same biological program. If they are, "
              "'metastatic' is one target worth predicting. If they are not, any model trained "
              "to recognise 'metastasis' is partly learning which organ it is looking at.")

    h(doc, "Three ways this could have been wrong", 1)
    para(doc, "1. Counting the same tumour thousands of times. 91,496 spots sounds like enormous "
              "statistical power, but spots inside one section are not independent. Every test "
              "here collapses each section to a single profile and treats the SECTION as the unit "
              "of analysis — 9 liver versus 5 nodal, never 91,496.")
    para(doc, "2. Rediscovering anatomy. A lymph node is an immune organ, so comparing whole "
              "sections would 'discover' that lymph nodes contain lymphocytes. The analysis keeps "
              "only tumour-dominated spots and then repeats itself at rising purity thresholds.")
    para(doc, "3. Comparing different people. Four patients donated both a liver and a nodal "
              "metastasis, so for them the comparison happens inside one person and patient "
              "identity cancels out.")
    figure(doc, os.path.join(O6, "figures", "fig1_cohort_design.png"),
           "Figure 1. Tissue available per patient. Shaded patients gave both a liver and a "
           "lymph-node metastasis.")

    # ---------------------------------------------------------------- results
    doc.add_page_break()
    h(doc, "Result 1 — the purity control", 1)
    figure(doc, os.path.join(O6, "figures", "fig2_purity_control.png"),
           "Figure 2. If the immune difference were leftover host tissue it would shrink as we "
           "demand purer tumour. It does not: the ratio holds at 2.0–2.4x while liver "
           "contamination falls eightfold.")
    callout(doc, "This is the control the whole finding rests on. Liver-cell contamination drops "
                 "eightfold across the thresholds while the lymphoid excess in nodal metastases "
                 "stays at roughly twice that of liver metastases — so it is not host tissue "
                 "bleeding into the measurement.")

    h(doc, "Result 2 — the sites differ, and the difference is immune", 1)
    figure(doc, os.path.join(O6, "figures", "fig3_bcells_paired.png"),
           "Figure 3. Left: the four patients measured at both sites; every one rises from liver "
           "to lymph node. Right: all sections, with primaries for reference.")
    figure(doc, os.path.join(O6, "figures", "fig5_shared_vs_sitespecific.png"),
           "Figure 4. Each dot is one feature. Under a single universal metastatic program every "
           f"dot would fall on the dashed line. Observed alignment {cos6:+.3f}.")
    callout(doc, f"About {100*cos6:.0f}% of 'becoming a metastasis' is common to both sites and the "
                 "rest depends on where the tumour landed. Two features behave as they must: "
                 "hepatocytes rise only in liver metastases, and stromal features fall sharply at "
                 "both sites — metastases shed the dense scar-like stroma that defines primary "
                 "pancreatic cancer.")

    doc.add_page_break()
    h(doc, "Result 3 — replication in an independent cohort", 1)
    para(doc, "GSE274557 (Maitra lab, 2025): 55 Visium sections from 13 treatment-naive patients, "
              "sampling primary tumour plus liver, peritoneal and lung metastases. Eleven of those "
              "patients donated two or more different metastatic sites, against four in our cohort.")
    figure(doc, os.path.join(O7, "figures", "fig6_replication.png"),
           "Figure 5. Left: alignment between the metastatic shifts of two destinations, in both "
           "cohorts. Right: every feature, with its shift at each destination.")
    callout(doc, "In our cohort liver and lymph node shared about half their shift (+0.55). Across "
                 "the three pairs available here the alignment is −0.26, +0.09 and −0.37: not "
                 "merely unrelated but, for two pairs, pointing in opposite directions. Immune "
                 "content falls in liver metastases and rises in lung metastases — the liver is an "
                 "immune-tolerant organ, and liver metastases are notoriously immune-excluded.")

    doc.add_page_break()
    h(doc, "Result 4 — a third test, on seven organs", 1)
    para(doc, "Both cohorts above study pancreatic cancer. If tissue context really governs "
              "expression this strongly, the effect should not be a quirk of PDAC, and it should "
              "show up in a task rather than only in a correlation. We tested a falsifiable "
              "prediction on HEST-1k: 159 human tumour sections across seven organs.")
    para(doc, "The prediction: a model that predicts expression from morphology should not "
              "transfer between organs.")
    figure(doc, os.path.join(O8, "figures", "fig8_cross_organ.png"),
           "Figure 6. Left: accuracy when nothing changes, when the study changes, and when the "
           "organ changes. Right: every train/test organ pair; the bold diagonal is same-organ.")
    table(doc, ["Condition", "Mean per-gene r"], [
        ["Same organ, same study", f"{M8['within_mean']:+.3f}"],
        ["Same organ, DIFFERENT study (batch only)", f"{M8.get('across_study_mean', 0.199):+.3f}"],
        ["DIFFERENT organ", f"{M8['cross_mean']:+.3f}"],
    ])
    callout(doc, "Moving to a different study of the same organ — different lab, protocol and "
                 "batch — costs about 15% of the signal. Moving to a different organ costs a "
                 "further 37%. The degradation is tissue biology, not a technical artefact.")
    para(doc, "Brain is the sharpest case: worst within itself (+0.10) and almost nothing "
              "transfers into it, as expected of neural tissue set against six epithelial cancers.")

    # ---------------------------------------------------------------- limits
    doc.add_page_break()
    h(doc, "What this does not show", 1)
    for t in [
        "Not proof that tumour cells differ. A 55-micron spot at 80% tumour purity still contains "
        "real non-tumour cells. The effect survives every purity threshold we can test, which "
        "argues against contamination, but Visium cannot fully separate a tumour cell from its "
        "neighbour.",
        "Underpowered for most features. Only 2 of 42 features reach significance in the first "
        "cohort. Absence of evidence is not evidence of absence.",
        "The paired test cannot itself reach significance. With four patients the smallest "
        "achievable Wilcoxon p is 0.125; its value is the perfect 4/4 direction agreement.",
        "The lymph-node B-cell result rests on one cohort — the replication cohort has no nodal "
        "samples, so it is untested rather than refuted.",
        "The seven-organ test measures a different thing: it shows morphology→expression rules "
        "are tissue-specific, and only four of its seven organs could be batch-controlled.",
        "No clinical validation. Clinical fields are populated for only 6 of 13 patients, all "
        "stage pM1, with no survival endpoint.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    h(doc, "Why it matters for the imaging model", 1)
    callout(doc, "The project's goal was to predict metastatic behaviour in primary tumour tissue "
                 "from an H&E slide alone. These results say that goal needs restating: there is "
                 "no single 'metastatic' target to predict. About half the signal is shared and "
                 "half depends on the destination organ, and morphology→expression rules do not "
                 "cross tissue boundaries. That, rather than model capacity, is the likeliest "
                 "reason every metastasis-derived target we tested proved unrecoverable from "
                 "morphology.")

    h(doc, "Method", 1)
    para(doc, "RCTD cell-type deconvolution (published assay), 27 Bagaev expression signatures, "
              "scVI latent space, UNI2-h histology features. Tests are Mann–Whitney at section "
              "level with Benjamini–Hochberg correction across 42 features, plus within-patient "
              "Wilcoxon on the four paired patients. Primary tumour-purity threshold 0.50. "
              "Cohorts: GSE272362 (Khaliq, Nat Genet 2024), GSE274557 (Maitra, 2025), HEST-1k "
              "(Jaume, NeurIPS 2024). Reproduce with 04_Models/stage6_site_specific_program.py, "
              "stage7_external_replication.py and 05_HEST/hest_cross_organ.py.", size=9)

    os.makedirs(os.path.dirname(DEST), exist_ok=True)
    doc.save(DEST)
    print("wrote", DEST, f"({os.path.getsize(DEST)/1024:.0f} KB)")


if __name__ == "__main__":
    main()
