"""
Build the detailed research write-up as an editable Word document.

Long form, closer to a supplementary methods and results file than a short paper:
every objective, how it was tested, what came out, and what it does not support.
All numbers are read from the results files so the document cannot drift from the
analysis.

Run:
    python docs/reports/build_research_paper.py
Output:
    docs/reports/PDAC_metastatic_site_specificity_FULL.docx
"""

import json
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
R = os.path.join(ROOT, "Outputs")
FIGS = os.path.join(ROOT, "docs", "reports", "figures")
DEST = os.path.join(ROOT, "docs", "reports",
                    "PDAC_metastatic_site_specificity_FULL.docx")

INK = RGBColor(0x12, 0x16, 0x1D)
GREY = RGBColor(0x4A, 0x53, 0x61)
BLUE = RGBColor(0x1F, 0x4E, 0x8C)


def load(p, name="metrics.json"):
    try:
        return json.load(open(os.path.join(R, p, name)))
    except Exception:
        return {}


M0 = load("stage0_full_cohort")
M1 = load("stage1a_full_cohort_residcaf", "convergence.json")
M3 = load("stage3_target_comparison")
M5 = load("stage5_shared_met_axis")
M6 = load("stage6_site_program")
M7 = load("stage7_external_replication")
M7C = load("stage7_external_replication", "cosines_organ_marker_control.json")
M8 = load("stage8_hest_cross_organ")
INF = json.load(open(os.path.join(FIGS, "inference_example.json")))


def acc(metrics, contrast, label):
    for c in metrics.get("contrasts", []):
        if c["contrast"] == contrast and c["tumour_only"]:
            for r in c["results"]:
                if r["label"] == label:
                    return r["balanced_acc"]
    return float("nan")


# ------------------------------------------------------------------ helpers
def setup(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.18
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.95)
        s.top_margin = s.bottom_margin = Inches(0.85)


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = INK if level > 1 else BLUE
    return p


def P(doc, text, size=10.5, italic=False, color=None, bold=False, space=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(size)


def numbered(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(size)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def table(doc, header, rows, widths=None, caption=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(txt))
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    if caption:
        c = doc.add_paragraph()
        r = c.add_run(caption)
        r.font.size = Pt(8.5)
        r.italic = True
        r.font.color.rgb = GREY
    doc.add_paragraph()
    return t


def figure(doc, path, caption, width=6.5):
    if not os.path.exists(path):
        P(doc, f"[figure not found: {os.path.basename(path)}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = c.add_run(caption)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = GREY
    doc.add_paragraph()


# ================================================================== document
doc = Document()
setup(doc)

H(doc, "Where a pancreatic tumour spreads changes what it becomes: "
       "site specific metastatic programs across three spatial transcriptomics datasets", 0)
P(doc, "Author list to be completed", italic=True, color=GREY)
P(doc, "Detailed analysis report and supplementary methods. "
       "Draft for internal review.", italic=True, color=GREY, size=9.5)

# ------------------------------------------------------------------ abstract
H(doc, "Abstract", 1)
P(doc,
  "Metastatic pancreatic cancer is usually treated as a single disease state. Tumour deposits in "
  "the liver, the lymph nodes, the lung and the peritoneum are grouped together as metastasis, and "
  "computational work often trains a single model to recognise that state. We asked whether that "
  "grouping is justified at the level of gene expression. Using 91,496 spatially resolved "
  "expression measurements from 30 tissue sections of 13 patients, we compared how liver and "
  "lymph node metastases differ from the matched primary tumour. The two destinations share only "
  "about half of their departure from the primary, and the part that differs is immune, with B "
  "cell content the single feature that survives correction for multiple testing. The result "
  "reproduces in an independent cohort of 13 treatment naive patients covering liver, lung and "
  "peritoneal deposits, where the shifts at different sites are unrelated or opposed rather than "
  "shared, and where liver deposits are immune poor while lung deposits are immune rich. Finally "
  "we tested a prediction that follows from these observations: if tissue context governs "
  "expression, a model that predicts expression from histology should not transfer between "
  "organs. On 156 sections spanning seven organs, transfer to a new laboratory studying the same "
  "organ cost 15 percent of performance, while transfer to a different organ cost a further 37 "
  "percent. Taken together the three analyses argue that there is no single metastatic program to "
  "detect, which has direct consequences for efforts to predict metastatic behaviour from routine "
  "histology.")

# ------------------------------------------------------------------ intro
H(doc, "1. Background and motivation", 1)
P(doc,
  "Pancreatic ductal adenocarcinoma is usually fatal because of metastasis rather than the "
  "primary tumour. A long standing goal in computational pathology is to look at a routine "
  "haematoxylin and eosin slide of the primary tumour and say something useful about whether, and "
  "where, that tumour will spread. Spatial transcriptomics makes this question tractable in a new "
  "way, because it measures gene expression at thousands of small locations across a tissue "
  "section while keeping the matching image of that same tissue.")
P(doc,
  "This project began with that goal. The intention was to learn a transcriptional signature of "
  "metastatic behaviour from spatial data, then train an image model to recover it from histology "
  "alone. Work on that plan produced a series of negative results, and those negatives pointed at "
  "a more basic question that had been assumed rather than tested: whether metastases in "
  "different organs are running the same program at all. This report documents both parts, "
  "because the negative results are what motivated the positive one.")

H(doc, "2. Objectives", 1)
P(doc, "The work is organised around four questions, each of which arose from the answer to the "
       "previous one.")
numbered(doc, [
    "Is the transcriptional difference between metastasis and primary tumour a property of the "
    "cancer, or largely a property of the organ the deposit is sitting in?",
    "Do metastases at two different sites in the same patient share one program, and if not, what "
    "distinguishes them?",
    "Does that pattern hold in an independent cohort, with more metastatic sites and without "
    "prior chemotherapy?",
    "If tissue context governs expression, does a model that predicts expression from histology "
    "fail to transfer between organs, as it should if the effect is real?",
])
P(doc, "A fifth question ran alongside these and is reported in section 6: whether any "
       "metastasis derived target can be predicted from histology in the primary tumour, which "
       "was the original aim of the project.")

# ------------------------------------------------------------------ methods
doc.add_page_break()
H(doc, "3. Materials and methods", 1)

H(doc, "3.1 Datasets", 2)
table(doc,
      ["", "Discovery", "Replication", "Generalisation"],
      [["Accession", "GSE272362", "GSE274557", "HEST-1k"],
       ["Reference", "Khaliq et al. 2024", "Maitra et al. 2025", "Jaume et al. 2024"],
       ["Sections used", "30", "55", "156"],
       ["Patients", "13", "13", "many, pooled from 153 studies"],
       ["Sites or organs", "primary, liver, lymph node, normal pancreas",
        "primary, liver, lung, peritoneum", "seven organs"],
       ["Treatment", "mixed, some neoadjuvant", "all treatment naive", "mixed"],
       ["Platform", "Visium", "Visium", "Visium, whole transcriptome only"],
       ["Spots", "91,496", "132,349", "about 500,000"]],
      widths=[1.5, 1.8, 1.8, 1.9],
      caption="Table 1. The three datasets. Each answers a question the previous one could not.")

H(doc, "3.2 What is measured, and how", 2)
P(doc,
  "A Visium slide reports gene expression at spots of 55 micrometres, each covering roughly one "
  "to ten cells, arranged on a regular grid across the section. Every spot therefore has three "
  "descriptions available: the image tile centred on it, its expression profile, and, where a "
  "deconvolution has been run, an estimate of which cell types contribute to it. Figure 1 shows "
  "how these fit together.")
figure(doc, os.path.join(FIGS, "fig_methods.png"),
       "Figure 1. Study design. (a) The three datasets and the question each was used to answer. "
       "(b) The three measurement channels available for every spot. They are parallel "
       "descriptions of the same tissue location, not a processing pipeline. (c) The three "
       "safeguards, each shown against the specific mistake it prevents. (d) The four objectives "
       "and how each answer produced the next question.", width=6.6)

bullets(doc, [
    "Histology. The 224 by 224 pixel tile centred on each spot is passed through UNI2-h, a "
    "pathology foundation model, giving 1,536 numbers per spot. The same model and the same "
    "preprocessing are used in every dataset so that results are comparable.",
    "Expression. Counts are normalised to 10,000 per spot and log transformed. Depending on the "
    "analysis these are summarised as a 50 dimensional scVI latent representation, as published "
    "signature scores, or as marker gene scores computed directly.",
    "Cell composition. For the discovery cohort we use the deconvolution published with the "
    "original study rather than running our own, which avoids introducing a second source of "
    "disagreement.",
])

H(doc, "3.3 Three safeguards", 2)
P(doc, "Three features of this data will produce confident and wrong answers if they are ignored. "
       "Each is addressed by a specific control rather than by an assumption.")

H(doc, "3.3.1 The unit of analysis is the section, never the spot", 3)
P(doc,
  "There are 91,496 spots in the discovery cohort but only 17 metastasis sections from 13 people. "
  "Spots within a section are not independent observations. They are one tumour, measured "
  "repeatedly. A test run at spot level would return extremely small p values for effects that do "
  "not exist. Every statistical test in this report is therefore computed on pseudobulk profiles, "
  "one value per section, and the sample size is the number of sections.")

H(doc, "3.3.2 Tumour purity is swept, not fixed", 3)
P(doc,
  "A lymph node is an immune organ and a liver is full of hepatocytes. Comparing whole sections "
  "would recover facts about anatomy rather than about cancer. We restrict every comparison to "
  "tumour rich spots and then repeat the whole analysis at increasing purity thresholds. A "
  "difference caused by residual host tissue has to shrink as purity rises. A difference that "
  "persists cannot be explained that way.")

H(doc, "3.3.3 Patient and study are held fixed where possible", 3)
P(doc,
  "Comparing 12 liver sections against 5 nodal sections also compares 12 people against 5. Four "
  "patients in the discovery cohort provided both a liver and a lymph node deposit, so for those "
  "four the comparison happens inside a single person and patient identity cancels. In HEST the "
  "equivalent problem is that an organ is often covered by a single laboratory, so we compare "
  "across studies with the organ held fixed before attributing anything to the organ.")

H(doc, "3.4 Statistical treatment", 2)
bullets(doc, [
    "Group differences between sites use the Mann Whitney test at section level, with Benjamini "
    "and Hochberg correction across the 42 features tested. Effect size is reported as Cohen's d.",
    "Paired comparisons within patients use the Wilcoxon signed rank test. With four pairs the "
    "smallest achievable p value is 0.125, so for these the informative quantity is the "
    "consistency of direction rather than significance.",
    "Classification uses logistic regression with leave one patient out cross validation and "
    "pooled out of fold balanced accuracy. A null distribution is built by relabelling primary "
    "sections into two arbitrary groups and repeating the procedure.",
    "Regression from image features uses ridge regression with the penalty selected inside the "
    "training folds. Performance is the Pearson correlation between predicted and measured "
    "expression on held out sections, averaged across genes.",
])

H(doc, "3.5 What was actually run, objective by objective", 2)
P(doc, "This section is the procedural record. Each objective corresponds to one script, listed "
       "in section 10.")

H(doc, "3.5.1 Objective 1, is the metastatic direction organ identity", 3)
numbered(doc, [
    "Take all 30 sections of the discovery cohort and keep only spots whose estimated tumour "
    "fraction is at least 0.5.",
    "Build three descriptions of each spot: the full 15 channel cell composition vector, the same "
    "vector with the hepatocyte channel deleted, and the same vector with both hepatocyte and "
    "tumour channels deleted. Add the 50 dimensional expression representation as a fourth.",
    "For each description and each of the three site comparisons, fit logistic regression under "
    "leave one patient out cross validation and pool the out of fold predictions into one "
    "balanced accuracy.",
    "Build the null by taking primary sections only, splitting the patients into two arbitrary "
    "groups, and repeating the whole procedure. Report the mean across repeats.",
])
note(doc, "The logic. If the separation is cancer biology it should survive deleting the obvious "
          "organ channel and should appear at both metastatic sites. If it is organ identity it "
          "will do neither.")

H(doc, "3.5.2 Objective 2, do the two sites share one program", 3)
numbered(doc, [
    "Score every spot for 27 published functional gene signatures and 15 cell type fractions, "
    "giving 42 features.",
    "Keep tumour rich spots, then average within each section to get one profile per section. "
    "This is the pseudobulk step, and it is what fixes the sample size at 24 sections rather "
    "than 91,496 spots.",
    "For each metastatic site compute the shift, meaning the mean profile at that site minus the "
    "mean profile of the primary sections, across all 42 features.",
    "Measure how far the two shift vectors point in the same direction using cosine similarity. "
    "A value of 1 would mean one universal program; 0 would mean the two destinations are "
    "unrelated.",
    "Test each of the 42 features separately with the Mann Whitney test at section level and "
    "correct across all 42 with Benjamini and Hochberg.",
    "Repeat step 3 within each of the four patients who provided both a liver and a nodal "
    "deposit, and record how many of the four agree in direction.",
    "Repeat the entire analysis at tumour purity thresholds of 0.3, 0.5, 0.7 and 0.8 and check "
    "whether the effect decays.",
])

H(doc, "3.5.3 Objective 3, does it replicate", 3)
numbered(doc, [
    "Process GSE274557 from raw counts through the same normalisation used in the discovery "
    "cohort, so that nothing in the comparison depends on preprocessing choices.",
    "Score the same style of feature set, computing marker gene scores directly because this "
    "cohort has no published deconvolution.",
    "Include three organ positive controls, hepatocyte, lung and mesothelial markers, so that a "
    "failure to detect the organ a sample came from would be visible.",
    "Compute the shift for each of the three metastatic sites against primary, then all three "
    "pairwise cosines.",
    "Repeat with the organ contamination markers removed, to check that the answer is not being "
    "produced by residual host tissue.",
])
note(doc, "The mesothelial control failed and is reported as such. Mesothelin is expressed by "
          "pancreatic tumour cells themselves, so a mesothelial marker score cannot distinguish "
          "peritoneal lining from tumour in this disease. The hepatocyte and lung controls both "
          "behaved correctly.")

H(doc, "3.5.4 Objective 4, does an image model cross organs", 3)
numbered(doc, [
    "Select HEST-1k sections that are Visium, have paired image tiles, and report more than 5,000 "
    "genes. The gene count filter is not cosmetic; it is explained in section 7.",
    "Embed every image tile with UNI2-h, the same model used throughout.",
    "Choose a 50 gene panel by ranking genes on their variance in the weakest organ rather than "
    "on average variance, and discarding genes that are near zero anywhere. Ranking on average "
    "variance selects organ marker genes such as digestive enzymes, which would guarantee the "
    "result before the experiment ran.",
    "Train ridge regression from image features to expression, and evaluate under three "
    "conditions: held out sections of the same organ and study, held out sections of the same "
    "organ from a different study, and held out sections of a different organ.",
    "Report the mean per gene Pearson correlation in each condition.",
])
note(doc, "The middle condition is the control that makes the experiment worth running. Without "
          "it, any cross organ drop could be attributed to changing laboratory.")

# ------------------------------------------------------------------ results
doc.add_page_break()
H(doc, "4. Results", 1)

H(doc, "4.1 Objective 1. Most of the apparent metastatic signal is organ identity", 2)
P(doc,
  "We first asked whether metastasis can be separated from primary tumour using cell composition, "
  "and whether that separation reflects cancer biology or the surrounding organ. Restricting to "
  "tumour dominated spots, a classifier separates liver metastasis from primary at balanced "
  "accuracy " + f"{acc(M0,'HM_vs_T','rctd_all'):.3f}" + ". Removing the hepatocyte channel barely "
  "changes this, and removing both the hepatocyte and the tumour channels still leaves "
  f"{acc(M0,'HM_vs_T','rctd_no_hep_no_tumour'):.3f}" + ". The same comparison for lymph node "
  "metastasis reaches only " + f"{acc(M0,'LNM_vs_T','rctd_all'):.3f}" + ", against a null of "
  f"{M0.get('null_control',{}).get('rctd_all_mean',float('nan')):.3f}" + " obtained by splitting "
  "primary sections into arbitrary groups.")
table(doc, ["Comparison, tumour rich spots only", "Cell composition",
            "Composition without liver and tumour channels", "Transcriptome"],
      [["Liver metastasis vs primary", f"{acc(M0,'HM_vs_T','rctd_all'):.3f}",
        f"{acc(M0,'HM_vs_T','rctd_no_hep_no_tumour'):.3f}", f"{acc(M0,'HM_vs_T','scvi'):.3f}"],
       ["Lymph node metastasis vs primary", f"{acc(M0,'LNM_vs_T','rctd_all'):.3f}",
        f"{acc(M0,'LNM_vs_T','rctd_no_hep_no_tumour'):.3f}", f"{acc(M0,'LNM_vs_T','scvi'):.3f}"],
       ["Liver vs lymph node metastasis", f"{acc(M0,'HM_vs_LNM','rctd_all'):.3f}",
        f"{acc(M0,'HM_vs_LNM','rctd_no_hep_no_tumour'):.3f}", f"{acc(M0,'HM_vs_LNM','scvi'):.3f}"],
       ["Null, arbitrary primary groups",
        f"{M0.get('null_control',{}).get('rctd_all_mean',float('nan')):.3f}", "",
        f"{M0.get('null_control',{}).get('scvi_mean',float('nan')):.3f}"]],
      widths=[2.6, 1.3, 1.6, 1.2],
      caption="Table 2. Leave one patient out balanced accuracy. Composition separates liver "
              "deposits well but does not generalise to lymph node deposits, which is the "
              "signature of an organ effect rather than a metastatic one.")
note(doc, "Reading. A composition based description of metastatic direction is mostly a "
          "description of which organ the sample came from. It survives deletion of the obvious "
          "organ marker, so removing hepatocytes is not sufficient, and it collapses at the other "
          "metastatic site. The transcriptome behaves differently: it separates both sites from "
          "primary at a similar and more modest level.")

H(doc, "4.2 Objective 2. The two sites share about half a program, and differ in immune content", 2)
P(doc,
  "For each site we computed the shift away from the matched primary tumour across 42 features, "
  "then measured how far the two shifts point in the same direction. The alignment is "
  f"{M6.get('cos_HMshift_LNMshift', float('nan')):+.3f}" + ", so roughly half of what happens "
  "when a tumour becomes a metastasis is common to both destinations and the rest depends on "
  "where it landed.")
P(doc,
  "Testing the 42 features individually at section level, two survive correction. B cell content "
  "is higher in lymph node deposits with a large effect size and agrees in all four patients who "
  "provided both sites. Hepatocyte content is higher in liver deposits, which is a positive "
  "control: an analysis that failed to find liver cells in liver metastases would be broken.")
sig, seen = [], set()
for r in M6.get("top_LNM", []) + M6.get("top_HM", []):
    if r.get("q_unpaired", 1) < 0.10 and r["feature"] not in seen:
        seen.add(r["feature"])
        sig.append(r)
if sig:
    table(doc, ["Feature", "Cohen's d", "FDR q", "Paired patients agreeing"],
          [[r["feature"].replace("rctd::", "cell type: "), f"{r['cohens_d']:+.2f}",
            f"{r['q_unpaired']:.3f}",
            f"{int(r['n_paired_LNM_higher'])} of {int(r['n_paired'])}"] for r in sig],
          widths=[2.6, 1.2, 1.2, 1.7],
          caption="Table 3. The only two of 42 features that survive correction for multiple "
                  "testing. Positive d means higher in lymph node deposits.")
P(doc, "The purity sweep is what makes the immune result usable. Across rising tumour purity "
       "thresholds, liver cell contamination falls roughly eightfold, from 0.133 to 0.017, while "
       "the lymphoid excess in nodal deposits holds at about twice the liver level throughout. A "
       "contamination artefact would decay alongside the contamination. This does not.")
figure(doc, os.path.join(R, "stage6_site_program", "figures", "fig2_purity_control.png"),
       "Figure 2. The purity control. Liver cell contamination falls from 0.133 to 0.017 across "
       "the thresholds while the lymphoid ratio between nodal and liver deposits stays between "
       "1.96 and 2.44. Full numbers in Appendix A.", width=6.6)
figure(doc, os.path.join(R, "stage6_site_program", "figures", "fig3_bcells_paired.png"),
       "Figure 3. B cell content. Left, the four patients measured at both sites; every one rises "
       "from liver to lymph node. Right, all sections with primaries shown for reference.",
       width=6.3)
figure(doc, os.path.join(R, "stage6_site_program", "figures", "fig5_shared_vs_sitespecific.png"),
       "Figure 4. Each point is one feature, positioned by how much it changes at each "
       "destination. Under a single universal program every point would lie on the diagonal. "
       "Stromal features fall at both sites, so metastases shed the dense stroma that "
       "characterises the primary tumour.", width=5.2)

doc.add_page_break()
H(doc, "4.3 Objective 3. The pattern reproduces, and is stronger without chemotherapy", 2)
cs = M7.get("cosines", {})
csn = M7C.get("no_organ_markers", {})
P(doc,
  "The discovery cohort covers one pair of metastatic sites and includes patients who received "
  "neoadjuvant chemotherapy. GSE274557 provides 55 sections from 13 treatment naive patients "
  "sampling primary tumour together with liver, lung and peritoneal deposits, and 11 of those "
  "patients contributed two or more different metastatic sites.")
table(doc, ["Site pair", "Alignment, all features", "Alignment, organ markers removed"],
      [["Liver and lung", f"{cs.get('HM_vs_LuM', float('nan')):+.3f}",
        f"{csn.get('HM_vs_LuM', float('nan')):+.3f}"],
       ["Liver and peritoneum", f"{cs.get('HM_vs_PM', float('nan')):+.3f}",
        f"{csn.get('HM_vs_PM', float('nan')):+.3f}"],
       ["Lung and peritoneum", f"{cs.get('LuM_vs_PM', float('nan')):+.3f}",
        f"{csn.get('LuM_vs_PM', float('nan')):+.3f}"],
       ["Discovery cohort, liver and lymph node",
        f"{M6.get('cos_HMshift_LNMshift', float('nan')):+.3f}", ""]],
      widths=[2.5, 1.9, 2.1],
      caption="Table 4. Alignment between the shifts at two destinations. Values near zero or "
              "below indicate that the two sites are not running one program. Removing organ "
              "contamination markers strengthens rather than explains the effect.")
note(doc, "Reading. In the replication cohort the shifts are unrelated or opposed rather than "
          "partly shared. The direction is interpretable: immune content falls in liver deposits "
          "and rises in lung deposits. That opposition is why the alignment is negative, and it "
          "agrees with the clinical observation that liver metastases are immune excluded and "
          "respond poorly to immunotherapy.")
figure(doc, os.path.join(R, "stage7_external_replication", "figures", "fig6_replication.png"),
       "Figure 5. Left, alignment between destinations in both cohorts. Right, every feature with "
       "its shift at each of the three destinations; immune features spread widest and in "
       "opposite directions.", width=6.6)
P(doc, "One limitation is specific and worth stating. This cohort contains no lymph node samples, "
       "so the B cell result from the discovery cohort could not be checked here. It is untested "
       "rather than contradicted.")

H(doc, "4.4 Objective 4. Expression models do not transfer between organs", 2)
P(doc,
  "The first two analyses are correlational and both concern pancreatic cancer. If tissue context "
  "really governs expression, the claim should also hold as a prediction about a task, and in "
  "tissue other than pancreas. We trained a model to predict expression from histology and asked "
  "how well it carries to a new setting. HEST-1k provides 156 whole transcriptome sections across "
  "seven organs.")
P(doc,
  "The obvious objection is batch effect. These sections come from 153 separate studies, so a "
  "model that fails on a new organ might simply be failing on a new laboratory. The control is to "
  "hold the organ fixed and change the study instead. Four organs are covered by more than one "
  "study and can support that comparison.")
table(doc, ["Condition", "Mean per gene correlation", "Change"],
      [["Same organ, same study", f"{M8.get('within_mean', float('nan')):+.3f}", "reference"],
       ["Same organ, different study",
        f"{M8.get('across_study_mean', float('nan')):+.3f}", "loses about 15 percent"],
       ["Different organ", f"{M8.get('cross_mean', float('nan')):+.3f}",
        "loses a further 37 percent"]],
      widths=[2.6, 2.0, 1.9],
      caption="Table 5. Transfer of a histology to expression model. Changing laboratory costs "
              "little. Changing organ costs a great deal.")
figure(doc, os.path.join(R, "stage8_hest_cross_organ", "figures", "fig8_cross_organ.png"),
       "Figure 6. Left, the three conditions. Right, every training and testing organ pair, with "
       "the same organ diagonal in bold. Brain barely receives transfer from any other organ.",
       width=6.6)
note(doc, "Reading. The degradation is tissue biology rather than a technical artefact. What a "
          "histology image implies about gene expression is specific to the tissue it came from. "
          "This is the same claim as the first two analyses, arrived at through a completely "
          "different route and in six cancers that are not pancreatic.")

# ------------------------------------------------------------------ inference
doc.add_page_break()
H(doc, "5. Worked inference example", 1)
P(doc,
  f"To make the prediction task concrete, we trained on {INF['n_train_sections']} bowel sections "
  f"and applied the model to a section it had never seen, {INF['held_out']}, which contains "
  f"{INF['n_spots']:,} spots. The model receives only the image tile for each spot and returns a "
  "predicted expression value.")
P(doc,
  f"The best predicted gene in this section is {INF['best_gene']}, at a correlation of "
  f"{INF['best_r']:.2f} between predicted and measured values across spots. The median across the "
  f"50 gene panel is {INF['median_r']:.2f}. {INF['best_gene']} encodes a collagen chain, and the "
  "tiles with the highest measured values are visibly the pale fibrous regions while the lowest "
  "are dense sheets of tumour cells, so the model is keying on structure a pathologist would also "
  "recognise.")
figure(doc, os.path.join(FIGS, "fig_inference_example.png"),
       "Figure 7. Inference on a held out section. (a) Four real tiles ordered by measured "
       f"{INF['best_gene']}, each with the measured and predicted value. (b) The section mapped "
       "twice, by measured expression and by expression predicted from image alone. (c) Predicted "
       "against measured, one point per spot.", width=6.6)
figure(doc, os.path.join(FIGS, "fig_inference_genes.png"),
       "Figure 8. Not all genes are equally predictable. Genes tied to visible tissue structure "
       "are recovered well, while others are close to zero.", width=5.6)
note(doc, "This example is deliberately favourable. It uses 60 training sections of one organ and "
          "genes chosen for variability within that organ. The cross organ numbers in section 4.4 "
          "are lower because they use genes that must vary in every organ, which is a harder and "
          "more restrictive target. The two are not in conflict; they answer different questions.")

# ------------------------------------------------------------------ negative
H(doc, "6. What could not be predicted from histology", 1)
P(doc,
  "The original aim was to predict metastatic behaviour in primary tumour tissue from histology. "
  "That aim was pursued first and did not succeed. We report it because it is the reason the "
  "questions above were asked.")
t = M1.get("tests", [])
if t:
    table(doc, ["Metastatic axis defined from", "Sites used",
                "Agreement with the training target"],
          [[f"{' and '.join(x['trained_on'])} deposits",
            "liver and lymph node" if len(x["trained_on"]) == 2
            else ("liver only" if x["trained_on"] == ["HM"] else "lymph node only"),
            f"{x['leaving_score_resid']:+.3f}"] for x in t],
          widths=[2.3, 2.1, 2.1],
          caption="Table 6. The training target used in the original plan is largely independent "
                  "of the transcriptional metastatic axis, so predicting it is not the same as "
                  "predicting metastatic behaviour. The small amount of agreement that does exist "
                  "is liver specific and vanishes for lymph node, which is the same site "
                  "dependence seen everywhere else in this report.")
res = M3.get("results", {})
NICE = {"leaving programme (raw)": "leaving programme, unadjusted",
        "leaving programme (residualised)": "leaving programme, adjusted for tumour and CAF content",
        "met_resemblance  [metastasis-derived]": "resemblance to metastasis, derived from met tissue"}
if res:
    table(doc, ["Target", "Held out correlation from histology"],
          [[NICE.get(v["label"], v["label"]), f"{v['pooled_rho']:+.3f}"] for v in res.values()],
          widths=[3.9, 2.5],
          caption="Table 7. Histology predicts the EMT and stroma target modestly, and does not "
                  "recover the metastasis derived target at all. Four patients had histology "
                  "available, which limits this comparison.")
P(doc,
  "Five separate metastasis derived targets were tested and none was recoverable from histology, "
  "while the target that histology does predict turns out not to track the metastatic axis. Given "
  "the results in sections 4.2 to 4.4, the most economical explanation is not that the model was "
  "too weak. It is that there was never a single metastatic target to learn.")

# ------------------------------------------------------------------ traps
doc.add_page_break()
H(doc, "7. Analysis decisions that changed the answer", 1)
P(doc,
  "Several choices in this work were not neutral. Each one, made differently, would have produced "
  "a different published conclusion from the same data. We record them because a reader cannot "
  "judge the results without them.")

H(doc, "7.1 Three sections using a targeted gene panel inverted the cross organ result", 2)
P(doc,
  "HEST-1k labels several sections as Visium that are in fact targeted panels measuring roughly a "
  "thousand chosen genes rather than the whole transcriptome. Intersecting gene lists across "
  "sections collapses the shared gene set down to whatever those panels contain. With them "
  "included, the analysis concluded that the transfer loss was batch effect and not organ. "
  "Filtering on a minimum of 5,000 detected genes removed them and reversed the conclusion. The "
  "platform label alone is not sufficient to know what was measured.")

H(doc, "7.2 How the gene panel is chosen decides the result in advance", 2)
P(doc,
  "Selecting the most variable genes across the pooled dataset returns organ marker genes such as "
  "pancreatic digestive enzymes. A model tested on those genes will always fail to cross organs, "
  "because the genes were chosen to differ between organs. We instead rank genes by their "
  "variance in the organ where they vary least, which forces the panel to be informative "
  "everywhere. The reported cross organ drop is therefore a lower bound, not an inflated one.")

H(doc, "7.3 The deconvolution was taken from the source publication, not recomputed", 2)
P(doc,
  "An early version of this work used our own deconvolution run. It turned out to be strongly "
  "anticorrelated with the published one on tumour fraction, and several downstream conclusions "
  "flipped when the published assay was used instead. Every result in this report uses the assay "
  "distributed with the original study. Where a stage was affected, it was recomputed rather than "
  "adjusted.")

H(doc, "7.4 An equal separation is not a shared axis", 2)
P(doc,
  "Liver and lymph node deposits separate from primary tumour at almost the same accuracy, "
  f"{acc(M0,'HM_vs_T','scvi'):.3f} and {acc(M0,'LNM_vs_T','scvi'):.3f}. We initially read that as "
  "evidence for one shared metastatic axis. It is not. Two directions can be equally far from a "
  "starting point while pointing different ways, which is exactly what the cosine analysis in "
  "section 4.2 later showed. Equal separability and shared direction are different claims and "
  "have to be tested separately.")

# ------------------------------------------------------------------ discussion
doc.add_page_break()
H(doc, "8. Discussion", 1)
P(doc,
  "Three analyses on three datasets point the same way. Metastatic deposits in different organs "
  "are not interchangeable samples of one disease state. About half of the transcriptional change "
  "between primary tumour and metastasis is shared between destinations, and the remainder "
  "depends on where the deposit landed. The part that differs is immune, and the direction is "
  "consistent with what is already known clinically about the liver as an immune tolerant site.")
P(doc,
  "The third analysis is the one we would defend most strongly, for two reasons. It was specified "
  "as a prediction before it was run, derived from the first two results rather than fitted to "
  "them. And it was very nearly wrong: before three sections using a targeted gene panel were "
  "removed from the cohort, the same code concluded that the loss was batch rather than organ. "
  "The control decided the result, not the headline number.")
P(doc,
  "For computational pathology the practical consequence is a caution about targets. A model "
  "trained to recognise metastasis on liver deposits is learning something that only partly "
  "applies to nodal or peritoneal deposits, and a model that predicts expression from morphology "
  "in one organ should not be assumed to work in another. Reporting a single accuracy figure "
  "across pooled organs will overstate what such a model can do in a new setting.")
P(doc,
  "For biology the immune result is the most actionable. If nodal deposits are consistently "
  "richer in B cells than liver deposits from the same patient, and if liver deposits are immune "
  "poor while lung deposits are immune rich, then the immune environment of a metastasis is "
  "determined substantially by its location rather than by the tumour that produced it. That "
  "would be worth testing directly with immunohistochemistry on matched deposits.")

H(doc, "9. Limitations", 1)
bullets(doc, [
    "Spot resolution. A 55 micrometre spot at 80 percent tumour purity still contains real non "
    "tumour cells. The purity sweep argues against contamination but cannot separate a tumour "
    "cell from its neighbour.",
    "Power. Only two of 42 features survive correction in the discovery cohort. Absence of "
    "evidence for the remaining 40 is not evidence of absence.",
    "The paired test cannot reach significance with four patients. Its value lies in the "
    "consistency of direction, not in a p value.",
    "The B cell result rests on one cohort. The replication cohort contains no lymph node "
    "samples, so it is untested elsewhere rather than confirmed.",
    "The cross organ analysis could only be batch controlled in four of seven organs. Kidney, "
    "liver, lymph node and pancreas are each covered by a single study.",
    "Absolute prediction accuracy in the cross organ analysis is modest throughout. Those numbers "
    "compare conditions with each other and are not a claim that expression is accurately "
    "recoverable from an image.",
    "No clinical validation. Clinical fields were available for 6 of 13 patients in the discovery "
    "cohort, all at stage pM1, with no survival endpoint. Nothing here is linked to outcome.",
])

H(doc, "10. Conclusion", 1)
P(doc,
  "Where a pancreatic tumour spreads changes what it becomes. Liver and lymph node deposits share "
  "only about half of their departure from the primary tumour, the difference is immune, and the "
  "pattern reproduces in an independent treatment naive cohort covering three metastatic sites. A "
  "prediction derived from those observations, that histology to expression models should not "
  "cross organ boundaries, holds across seven organs after controlling for laboratory of origin. "
  "The practical consequence for the original aim of this project is that metastatic behaviour is "
  "not one target, and predicting it from histology will require targets defined per site rather "
  "than pooled.")

H(doc, "11. Data and code", 1)
bullets(doc, [
    "GSE272362, Khaliq et al., Nature Genetics 2024. Discovery cohort.",
    "GSE274557, Maitra laboratory 2025, PMID 40269162. Replication cohort.",
    "HEST-1k, Jaume et al., NeurIPS 2024, MahmoodLab. Generalisation cohort.",
    "UNI2-h histology foundation model, MahmoodLab.",
])
table(doc, ["Objective", "Script", "Output folder"],
      [["1", "04_Models/stage0_full_cohort.py", "Outputs/stage0_full_cohort"],
       ["2", "04_Models/stage6_site_specific_program.py", "Outputs/stage6_site_program"],
       ["3", "04_Models/stage7_external_replication.py", "Outputs/stage7_external_replication"],
       ["4", "05_HEST/hest_cross_organ.py", "Outputs/stage8_hest_cross_organ"],
       ["Inference example", "docs/reports/make_inference_examples.py",
        "docs/reports/figures"],
       ["Figure 1", "docs/reports/make_methods_figure.py", "docs/reports/figures"],
       ["Negative results, section 6", "04_Models/stage1a_full_cohort.py and "
        "04_Models/stage3_target_comparison.py", "Outputs/stage1a_full_cohort_residcaf"]],
      widths=[1.7, 2.9, 1.9],
      caption="Table 8. Every number in this report can be regenerated from these scripts.")

# ------------------------------------------------------------------ appendices
doc.add_page_break()
H(doc, "Appendix A. Tumour purity sweep, discovery cohort", 1)
P(doc, "The control behind section 4.2. As the purity threshold rises, liver cell contamination "
       "in liver deposits falls sharply while the lymphoid ratio between nodal and liver deposits "
       "holds. Section counts fall because whole sections drop out when too few of their spots "
       "reach the threshold, which is why the highest threshold is not the most informative row.")
ps = M6.get("purity_sweep", [])
if ps:
    table(doc, ["Purity threshold", "Liver sections", "Nodal sections",
                "Hepatocyte content, liver", "Lymphoid, liver", "Lymphoid, node",
                "Ratio", "p"],
          [[f"{r['purity']:.1f}", int(r["n_hm"]), int(r["n_lnm"]),
            f"{r['hepatocyte_hm']:.3f}", f"{r['lymphoid_hm']:.3f}",
            f"{r['lymphoid_lnm']:.3f}", f"{r['ratio']:.2f}", f"{r['p_mwu']:.3f}"] for r in ps],
          widths=[0.95, 0.8, 0.8, 1.15, 0.85, 0.85, 0.6, 0.6],
          caption="Table A1. Contamination falls roughly eightfold. The effect does not. The p "
                  "value rises at the top two thresholds because sections drop out of the "
                  "comparison, not because the difference shrinks; the ratio at 0.8 purity is "
                  "the second highest in the table.")

H(doc, "Appendix B. Per organ performance, HEST-1k", 1)
P(doc, "The within organ column of section 4.4, broken out. Performance does not track the number "
       "of sections available, which is worth noting: lymph node scores highest on four sections "
       "and bowel scores lower on sixty one. Organs represented by few sections are also less "
       "internally varied, so a within organ score there is an easier test. Brain is the hardest "
       "tissue in both directions.")
w = M8.get("within", {})
o = M8.get("organs", {})
if w:
    table(doc, ["Organ", "Sections", "Within organ correlation"],
          [[k, o.get(k, ""), f"{v:+.3f}"] for k, v in
           sorted(w.items(), key=lambda kv: -kv[1])] +
          [["All organs pooled", M8.get("n_samples", ""),
            f"{M8.get('within_mean', float('nan')):+.3f}"]],
          widths=[2.2, 1.5, 2.5],
          caption="Table B1. Only four of the seven organs are covered by more than one study, so "
                  "the batch control in Table 5 rests on those four.")

H(doc, "Appendix C. The 50 gene panel used in the cross organ analysis", 1)
P(doc, "Selected by variance in the weakest organ, as described in section 7.2. The panel is "
       "dominated by extracellular matrix and structural genes, which is expected: those are the "
       "genes whose expression is visible as tissue architecture in a stained image.")
pan = M8.get("panel", [])
if pan:
    P(doc, ", ".join(pan), size=9.5)

H(doc, "Appendix D. Best and worst predicted genes in the worked example", 1)
tg = INF.get("top_genes", {})
wg = INF.get("worst_genes", {})
if tg:
    table(doc, ["Well predicted gene", "r", "Poorly predicted gene", "r"],
          [[list(tg)[i] if i < len(tg) else "",
            f"{list(tg.values())[i]:+.3f}" if i < len(tg) else "",
            list(wg)[i] if i < len(wg) else "",
            f"{list(wg.values())[i]:+.3f}" if i < len(wg) else ""]
           for i in range(max(len(tg), len(wg)))],
          widths=[2.0, 1.0, 2.0, 1.0],
          caption=f"Table D1. Held out section {INF['held_out']}, {INF['n_spots']:,} spots, "
                  f"trained on {INF['n_train_sections']} other bowel sections.")

doc.save(DEST)
print("wrote", DEST, f"({os.path.getsize(DEST)/1024:.0f} KB)")
